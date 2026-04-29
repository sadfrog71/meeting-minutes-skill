#!/usr/bin/env python3
"""
会议纪要 Word 文档生成脚本
根据结构化数据生成标准格式会议纪要
"""

import sys
import os
import argparse

# 依赖检查
try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误: 请先安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


# ========== 颜色定义 ==========
COLORS = {
    'primary_blue': RGBColor(0, 51, 102),        # 一级标题 - 深蓝
    'primary_cyan': RGBColor(0, 102, 153),       # 二级标题/强调 - 深青
    'border_gray': RGBColor(204, 204, 204),      # 边框 - 浅灰
    'white': RGBColor(255, 255, 255),
}


def rgb_to_hex(rgb_color):
    return str(rgb_color)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(
        '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), rgb_to_hex(color))
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, color=None, sz=4):
    """设置单元格边框"""
    if color is None:
        color = COLORS['border_gray']
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders {}>'.format(nsdecls('w')) +
        ''.join([
            r'<w:{} w:val="single" w:sz="{}" w:color="{}"/>'.format(
                edge, sz, rgb_to_hex(color)
            ) for edge in ['top', 'left', 'bottom', 'right']
        ]) + '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_table_cell_margins(cell):
    """设置单元格内边距"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        '<w:tcMar {}>'.format(nsdecls('w')) +
        r'<w:top w:w="80" w:type="dxa"/>' +
        r'<w:bottom w:w="80" w:type="dxa"/>' +
        r'<w:left w:w="100" w:type="dxa"/>' +
        r'<w:right w:w="100" w:type="dxa"/>' +
        r'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_header(doc, company_name, date_str):
    """添加文档头部"""
    # 公司名称 + 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name} {date_str}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS['primary_blue']
    run.font.name = '黑体'
    p.paragraph_format.space_after = Pt(6)


def add_title(doc, meeting_title):
    """添加会议标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meeting_title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COLORS['primary_blue']
    run.font.name = '黑体'
    p.paragraph_format.space_after = Pt(12)


def add_info_table(doc, info_data):
    """添加信息表格"""
    # info_data: [(label, value), ...]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(13)

    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        
        # 标签列
        label_cell = row.cells[0]
        label_cell.text = label
        set_cell_shading(label_cell, COLORS['primary_cyan'])
        set_cell_border(label_cell, COLORS['border_gray'])
        set_table_cell_margins(label_cell)
        for p in label_cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = COLORS['white']
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 值列
        value_cell = row.cells[1]
        value_cell.text = value
        set_cell_border(value_cell, COLORS['border_gray'])
        set_table_cell_margins(value_cell)
        for p in value_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_section_heading(doc, text, level=1):
    """添加章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(12)
        run.font.color.rgb = COLORS['primary_cyan']
        run.font.name = '黑体'
        # 下划线
        p.paragraph_format.border_bottom = None
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS['primary_blue']
        run.font.name = '黑体'

    return p


def add_bullet_point(doc, text, indent=0):
    """添加项目符号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.25)
    
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    
    return p


def add_numbered_point(doc, number, text):
    """添加编号段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(10.5)
    run_num.font.name = '宋体'
    
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.name = '宋体'
    
    return p


def add_table_with_header(doc, headers, rows, col_widths=None):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        set_cell_shading(cell, COLORS['primary_cyan'])
        set_cell_border(cell, COLORS['border_gray'], sz=6)
        set_table_cell_margins(cell)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = COLORS['white']
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '黑体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            set_cell_border(cell, COLORS['border_gray'], sz=4)
            set_table_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'

    return table


def add_attachment_note(doc, text="[见附件]"):
    """添加附件标注"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"附件：{text}")
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.italic = True


def generate_meeting_minutes(
    output_path,
    company_name="公司名称",
    date_str="YYYY年MM月DD日",
    meeting_title="会议纪要",
    time_str="",
    location="",
    theme="",
    attendees="",
    host="",
    absent="",
    recorder="",
    sections=None,
    attachments=None
):
    """
    生成会议纪要 Word 文档
    
    Args:
        output_path: 输出文件路径
        company_name: 公司名称
        date_str: 日期字符串
        meeting_title: 会议标题
        time_str: 时间
        location: 地点
        theme: 主题
        attendees: 参会人员
        host: 主持人
        absent: 请假
        recorder: 记录人
        sections: 章节内容列表，每项为 (标题, 内容列表)
                 内容列表每项为 (类型, 数据)
                 类型: 'heading1', 'heading2', 'bullet', 'numbered', 'table'
        attachments: 附件列表
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 头部
    add_header(doc, company_name, date_str)
    add_title(doc, meeting_title)

    # 信息表格
    info_data = [
        ("时间", time_str),
        ("地点", location),
        ("主题", theme),
        ("参会人员", attendees),
        ("主持", host),
        ("请假", absent),
        ("记录", recorder),
    ]
    add_info_table(doc, info_data)

    # 章节内容
    if sections:
        for section_title, section_content in sections:
            add_section_heading(doc, section_title, level=1)
            
            if isinstance(section_content, list):
                for item in section_content:
                    if isinstance(item, tuple):
                        item_type, item_data = item
                        if item_type == 'heading2':
                            add_section_heading(doc, item_data, level=2)
                        elif item_type == 'bullet':
                            add_bullet_point(doc, item_data)
                        elif item_type == 'numbered':
                            num, text = item_data
                            add_numbered_point(doc, num, text)
                        elif item_type == 'table':
                            headers, rows = item_data
                            add_table_with_header(doc, headers, rows)
                    elif isinstance(item, str):
                        add_bullet_point(doc, item)

    # 附件
    if attachments:
        for att in attachments:
            add_attachment_note(doc, att)
    else:
        add_attachment_note(doc, "无")

    # 保存
    doc.save(output_path)
    print(f"会议纪要已生成: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description='生成会议纪要 Word 文档')
    parser.add_argument('-o', '--output', required=True, help='输出文件路径')
    parser.add_argument('--company', default='公司名称', help='公司名称')
    parser.add_argument('--date', default='', help='日期')
    parser.add_argument('--title', default='会议纪要', help='会议标题')
    parser.add_argument('--time', default='', help='时间')
    parser.add_argument('--location', default='', help='地点')
    parser.add_argument('--theme', default='', help='主题')
    parser.add_argument('--attendees', default='', help='参会人员')
    parser.add_argument('--host', default='', help='主持人')
    parser.add_argument('--absent', default='', help='请假')
    parser.add_argument('--recorder', default='', help='记录人')
    
    args = parser.parse_args()

    # 生成文档
    generate_meeting_minutes(
        output_path=args.output,
        company_name=args.company,
        date_str=args.date,
        meeting_title=args.title,
        time_str=args.time,
        location=args.location,
        theme=args.theme,
        attendees=args.attendees,
        host=args.host,
        absent=args.absent,
        recorder=args.recorder,
    )


if __name__ == '__main__':
    main()
